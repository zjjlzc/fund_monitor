# -*- coding:utf-8 -*-  
"""
--------------------------------
    @Author: Dyson
    @Contact: Weaver1990@163.com
    @file: proxy_spider.py
    @time: 2017/3/9 16:27
--------------------------------
"""
import sys
import os

import pandas as pd
import scrapy
import fund_monitor.items
import re
import traceback
import datetime
import bs4
import time
import json
import numpy as np

log_path = r'%s/log/spider_DEBUG(%s).log' %(os.getcwd(),datetime.datetime.date(datetime.datetime.today()))

sys.path.append(sys.prefix + "\\Lib\\MyWheels")
sys.path.append(os.getcwd()) #########
reload(sys)
sys.setdefaultencoding('utf8')

import driver_manager
driver_manager = driver_manager.driver_manager()
import requests_manager
requests_manager = requests_manager.requests_manager()
import mysql_connecter
mysql_connecter = mysql_connecter.mysql_connecter()
import requests

from docx import Document

import set_log  # log_obj.debug(文本)  "\x1B[1;32;41m (文本)\x1B[0m"
log_obj = set_log.Logger(log_path, set_log.logging.WARNING, set_log.logging.DEBUG)
log_obj.cleanup(log_path, if_cleanup=True) # 是否需要在每次运行程序前清空Log文件

url_dict = {
    "ht_research":"海通研究",
    "essence_research":"安信证券研究",
    "gh_9d1fbd62976c":"长江研究",
    "guangfacelve":"广发策略研究",
    "ebresearch":"光大研究S",
    "gtjaresearch":"国泰君安证券研究",
    "gh_294e097afc75":"中信研究"
}

end_date = datetime.datetime.now()# - datetime.timedelta(days=1)# datetime.datetime(year=2018, month=1, day=10)
file_path = u"%s\%s微信文章.xlsx" %(os.getcwd(), end_date.strftime('%Y%m%d'))
# if os.path.exists(file_path):
#     os.remove(file_path)

class Spider(scrapy.Spider):
    name = "0013"

    def start_requests(self):
        # self.urls = ["http://weixin.sogou.com/weixin?type=1&s_from=input&query=%s&ie=utf8&_sug_=n&_sug_type_=" %key for key in url_dict]
        for key in url_dict:
            url = "http://weixin.sogou.com/weixin?type=1&s_from=input&query=%s&ie=utf8&_sug_=n&_sug_type_=" % key
            item = fund_monitor.items.FundMonitorItem()
            item['data'] = url_dict[key]
            yield scrapy.Request(url=url, meta={'item': item}, callback=self.parse)

    def parse(self, response):
        try:
            bs_obj = bs4.BeautifulSoup(response.text, 'html.parser')
            e_a = bs_obj.find('a', uigs='account_name_0')
            url = e_a.get('href')
            item = response.meta['item']
            yield scrapy.Request(url=url, meta={'item': item}, callback=self.parse1)
        except:
            log_obj.error("0013 parse %s 出错:\n%s" % (response.url, traceback.format_exc()))

    def parse1(self, response):
        try:
            item = response.meta['item']
            driver = driver_manager.initialization()
            print response.url
            driver.get(response.url)
            driver.get_screenshot_as_file('catlog.png')

            bs_obj = bs4.BeautifulSoup(driver.page_source, 'html.parser')
            while bs_obj.find('div', class_="weui_cell weui_vcode verifycode"):
                driver.get_screenshot_as_file('catlog.png')
                img_url = 'http://mp.weixin.qq.com' + bs_obj.find('img', id="verify_img").get('src')
                requests_manager.get_file(img_url, 'verify_img.png')

                verifycode = raw_input("输入验证码:")
                driver.find_element_by_class_name('weui_input frm_input').send_key(verifycode) # weui_btn weui_btn_primary btn
                driver.find_element_by_class_name('weui_btn_primary').click()
                time.sleep(2)

            bs_obj = bs4.BeautifulSoup(driver.page_source, 'html.parser')
            driver.quit()
            e_divs = bs_obj.find_all('div', class_='weui_media_bd')
            for e_div in e_divs:
                s = e_div.find('p', class_='weui_media_extra_info').get_text(strip=True)
                date = re.search(ur'\d{4}年\d{1,2}月\d{1,2}日', s).group()
                global end_date
                if datetime.datetime.strptime(date, ur'%Y年%m月%d日') < end_date:
                    break

                url = 'http://mp.weixin.qq.com' + e_div.find('h4', class_='weui_media_title').get('hrefs')
                title = e_div.find('h4', class_='weui_media_title').get_text(strip=True)
                company = item['data']

                # file_path = u"%s\%s微信文章.docx" %(os.getcwd(), datetime.datetime.now().strftime('%Y%m%d'))
                # if not os.path.exists(file_path):
                #     document = Document()
                #     document.save(file_path)
                #
                # document = Document(file_path)
                # document.add_paragraph(u"%s\n%s\n%s\n\n" %(company, title, url))
                # document.save(file_path)

                global file_path
                if not os.path.exists(file_path):
                    pd.DataFrame([]).to_excel(file_path)

                pd.read_excel(file_path).append([company, title, url, '_'], ignore_index=True).to_excel(file_path, index=None)
                print 'file saved:', file_path
                time.sleep(10)

                # driver = driver_manager.initialization()
                # driver.get(url)
                #
                # # bs_obj = bs4.BeautifulSoup(driver.page_source, 'html.parser')
                # driver.get_screenshot_as_file('page.png')
                # title = re.search(r'(?<=<h2 class="rich_media_title" id="activity-name">).+?(?=</h2>)', driver.page_source, re.S).group().strip()
                # date = re.search(r'(?<=<em id="post-date" class="rich_media_meta rich_media_meta_text">).+?(?=</em>)', driver.page_source, re.S).group()
                # company = item['data']# re.search(r'(?<=<span class="rich_media_meta rich_media_meta_text rich_media_meta_nickname">).+?(?=</em>)', driver.page_source, re.S).group()
                # content = re.search(r'(?<=id="js_content">).+?(?=</div>)', driver.page_source, re.S).group()
                # print (company,date,title)
                # with open(item['data'] + u".txt", 'a') as f:
                #     f.write("%s%s：%s\n" %(company,date,title))
                #     for s in re.findall(r'<p.*?>.+?</p.*?>', content):
                #         for s0 in re.findall(r'<.+?>', s):
                #             print u'正文去掉了：', s0
                #         s = re.sub(r'<.+?>', '', s)
                #         f.write(s + '\n')
                #     f.write('\n')
                #
                # driver.quit()
                # time.sleep(15)
        except:
            log_obj.error("0013 parse1 %s 出错:\n%s" %(response.url, traceback.format_exc()))






if __name__ == '__main__':
    pass